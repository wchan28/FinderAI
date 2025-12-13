from __future__ import annotations

from pathlib import Path
from typing import List, Dict, Union

from docx import Document


def extract_text_from_docx(file_path: Union[str, Path]) -> List[Dict]:
    """
    Extract text from a Word document.

    Returns a list with a single dict containing all document text:
    [
        {"page_number": 1, "text": "document content..."},
    ]

    Note: Word documents don't have strict page boundaries like PDFs,
    so we treat the entire document as one unit for simplicity.
    """
    file_path = Path(file_path)
    doc = Document(file_path)

    text_parts = []

    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        if para_text:
            text_parts.append(para_text)

    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    full_text = "\n".join(text_parts)

    if not full_text.strip():
        return []

    return [{"page_number": 1, "text": full_text}]


def extract_full_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extract all text from a Word document as a single string.
    """
    content = extract_text_from_docx(file_path)
    if content:
        return content[0]["text"]
    return ""
